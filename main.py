from kivy.app import App
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.widget import Widget
from kivy.uix.pagelayout import PageLayout
from kivy.properties import ObjectProperty
from kivy.uix.image import Image
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.graphics import Color, Rectangle, Line
from kivy.core.window import Window

from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.backends import default_backend
import DataBase
import os
import docx
import crypto
import socket,threading
import sys 

global s
s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
host = "127.0.0.1"
port = 5005

data = []




def hashing(string:str) -> str:
    digest = hashes.Hash(hashes.MD5(), backend=default_backend())
    digest.update(string.encode('utf-8'))
    dataHash = digest.finalize()
    return str(dataHash)

def time(sec:int) -> str:
        minute = sec / 60
        hours = int(minute / 60)
        return str(hours)

class SaveDialog(FloatLayout):
    save = ObjectProperty(None)
    text_input = ObjectProperty(None)
    cancel = ObjectProperty(None)

class LoadDialog(FloatLayout):
    load = ObjectProperty(None)
    cancel = ObjectProperty(None)
    
class Root(PageLayout):
    global s
    loadfile = ObjectProperty(None)
    savefile = ObjectProperty(None)
    text_input = ObjectProperty(None)
    file_name = ObjectProperty(None)
    time_create = ObjectProperty(None)
    size_file = ObjectProperty(None)
    check = ObjectProperty(None)
    text = ''
    showMessage = ObjectProperty(None)
    thread = threading.Thread()
    
    def __init__(self, **kwargs):
        super(Root, self).__init__(**kwargs)
        s.connect((host,port))
        self.thread = threading.Thread(target=self.handle_messages, daemon=True).start()

    
    def send_message(self,to_send_out):
        try:
            s.send((data[0]+" - "+to_send_out).encode('utf-8'))
        except Exception as e:
            print ("Error sending: ",e)

    def handle_messages(self):
        while True:
            try:
                message = s.recv(1024)
                message = message.decode('utf-8')
                self.showMessage.text += message + '\n'
            except Exception as e:
                print (e)

    def fileAuthenticity(self, author:str) -> None:
        if not author:
            pass
        else:
            data = author.split(' ')
            if DataBase.check(data[0], data[1]):
                self.check.active = True

    def crypt(self, text:str) -> None:
        self.text = crypto.encryption(self.text_input.text)
        self.text_input.text = self.text
        #self.text_input.text += '\n' + data[1] + data[2]
        #self.text_input.text = text

    def decrypt(self, text:str) -> None:
        self.text = crypto.decryption(self.text_input.text)
        self.text_input.text = self.text

    def dismiss_popup(self) -> None:
        self._popup.dismiss()

    def loadFile(self) -> None:
        content = LoadDialog(load=self.load, cancel=self.dismiss_popup)
        self._popup = Popup(title="Load file", content=content, size_hint=(1, 1))
        self._popup.open()

    def show_save(self) -> None:
        content = SaveDialog(save=self.save, cancel=self.dismiss_popup)
        self._popup = Popup(title="Save file", content=content, size_hint=(0.9, 0.9))
        self._popup.open()

    def load(self, path:str, filename:str) -> None:
        self.file_name.text = os.path.basename(filename[0])
        print(os.path.getctime(filename[0]))
        self.time_create.text = time(os.path.getctime(filename[0])) + ' часов назад'
        self.size_file.text = str(os.path.getsize(filename[0])) + ' byte'
        doc = docx.Document(os.path.join(path, filename[0]))
        text = ''
        self.fileAuthenticity(doc.core_properties.author)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        self.text_input.text = text
        self.dismiss_popup()

    def save(self, path:str, filename:str) -> None:
        document = docx.Document()
        document.add_paragraph(self.text)
        document.add_page_break()
        document.core_properties.author = data[1] + ' ' + data[2]
        document.save(os.path.join(path, filename))
        #with open(os.path.join(path, filename), 'w') as stream:
            #stream.write(str(self.text))
    
        self.dismiss_popup()

    
        
class MyApp(App):
    def build(self):
        self.title = 'Криптопрограмма'



class Check(FloatLayout):
    Window.size = (350, 250)
    def check_data(self, fullname:str, password:str) -> None:
        hashFullName = hashing(fullname)
        hashPassword = hashing(password)
        if DataBase.check(hashFullName, hashPassword):
            data.extend([fullname, hashFullName, hashPassword])
            AuthenticationApp().stop()
            MyApp().run()


   
class AuthenticationApp(App):
    def build(self):
        self.title = 'Аутентификация'



if __name__ == '__main__':
    AuthenticationApp().run()

sys.exit()